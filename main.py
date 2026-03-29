import os
import sys

# Fix Windows console encoding for emoji/unicode support
if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
    sys.stderr.reconfigure(encoding='utf-8', errors='replace')

# Python 3.13+ compatibility: audioop was removed, use audioop-lts
try:
    import audioop
except ImportError:
    import audioop_lts as audioop  # type: ignore
    sys.modules['audioop'] = audioop

import moviepy.editor as mp
from pydub import AudioSegment
from pydub.utils import which
from pydub.effects import normalize
from faster_whisper import WhisperModel
from transformers import AutoModelForSeq2SeqLM, AutoTokenizer
from gtts import gTTS
import nltk
from time import sleep, time
import torch
from langdetect import detect, LangDetectException, detect_langs
import json
import re
import unicodedata
from googletrans import Translator

# ✅ Ensure FFmpeg path is set (especially for Windows)
AudioSegment.converter = which("ffmpeg")

# Ensure NLTK tokenizer is available
nltk.download('punkt', quiet=True)
nltk.download('punkt_tab', quiet=True)

# Language code mapping for Whisper and other services
LANGUAGE_MAP = {
    'en': 'english',
    'es': 'spanish',
    'fr': 'french',
    'de': 'german',
    'it': 'italian',
    'pt': 'portuguese',
    'ru': 'russian',
    'ja': 'japanese',
    'zh': 'chinese',
    'hi': 'hindi',
    'ar': 'arabic',
    'ko': 'korean',
    'ta': 'tamil',
}

LANGUAGE_NAMES = {
    'en': 'English',
    'es': 'Spanish',
    'fr': 'French',
    'de': 'German',
    'it': 'Italian',
    'pt': 'Portuguese',
    'ru': 'Russian',
    'ja': 'Japanese',
    'zh': 'Chinese',
    'hi': 'Hindi',
    'ar': 'Arabic',
    'ko': 'Korean',
    'ta': 'Tamil',
}

# ===== SESSION-BASED PROGRESS & CANCEL TRACKING =====
# Per-session progress tracking (thread-safe)
import threading
_session_progress = {}  # {session_id: {"stage": "...", "percent": 0, "language": "..."}}
_session_cancel = set()  # Set of session_ids that requested cancel
_progress_lock = threading.Lock()
_cancel_lock = threading.Lock()

# Default progress for when no session is active
_default_progress = {"stage": "Idle", "percent": 0, "language": "English"}


def get_progress_status(session_id=None):
    """Get progress status for a specific session."""
    if session_id is None:
        return _default_progress.copy()
    with _progress_lock:
        return _session_progress.get(session_id, _default_progress.copy())


def set_progress_status(session_id, stage, percent, language="English"):
    """Set progress status for a specific session."""
    with _progress_lock:
        _session_progress[session_id] = {
            "stage": stage,
            "percent": percent,
            "language": language
        }


def request_cancel(session_id):
    """Request cancellation for a specific session."""
    if session_id:
        with _cancel_lock:
            _session_cancel.add(session_id)


def is_cancelled(session_id):
    """Check if a session has been cancelled."""
    if session_id is None:
        return False
    with _cancel_lock:
        return session_id in _session_cancel


def clear_session(session_id):
    """Clear session data after processing completes."""
    if session_id:
        with _progress_lock:
            _session_progress.pop(session_id, None)
        with _cancel_lock:
            _session_cancel.discard(session_id)


# Thread-safe model caching
_whisper_model = None
_summarizer = None
_model_lock = threading.Lock()

# Auto-detect device (GPU if available, else CPU)
def get_device():
    """Auto-detect best available device."""
    if torch.cuda.is_available():
        return "cuda"
    elif hasattr(torch.backends, 'mps') and torch.backends.mps.is_available():
        return "mps"  # Apple Silicon
    else:
        return "cpu"

def get_compute_type(device):
    """Get optimal compute type based on device."""
    if device == "cuda":
        return "float16"  # Faster on GPU
    elif device == "mps":
        return "float16"  # Apple Silicon
    else:
        return "int8"  # CPU optimization


# Thread-local storage for current session context
_current_session = threading.local()


def set_current_session(session_id):
    """Set the current session ID for this thread."""
    _current_session.id = session_id


def get_current_session():
    """Get the current session ID for this thread."""
    return getattr(_current_session, 'id', None)


def update_progress(stage, percent):
    """Helper to update progress stage and percentage for current session."""
    session_id = get_current_session()
    if session_id:
        set_progress_status(session_id, stage, percent)
    # Also update default progress for backward compatibility
    _default_progress["stage"] = stage
    _default_progress["percent"] = percent


def detect_language(text):
    """Detect language from text using multi-method approach for better accuracy."""
    try:
        # Method 1: langdetect with confidence scoring
        lang_probs = detect_langs(text)
        if lang_probs and len(lang_probs) > 0:
            best_lang = lang_probs[0]
            lang_code = best_lang.lang
            confidence = best_lang.prob
            print(f"\u2705 Language detection confidence: {confidence:.2%}")
        else:
            lang_code = detect(text)
            confidence = 1.0
        
        # Validate language code
        if lang_code not in LANGUAGE_NAMES:
            # Map similar languages
            if lang_code.startswith('zh'):
                lang_code = 'zh'
            elif lang_code.startswith('pt'):
                lang_code = 'pt'
            else:
                lang_code = 'en'
        
        lang_name = LANGUAGE_NAMES.get(lang_code, 'English')
        update_progress(f"\ud83c\udf10 Detected language: {lang_name} ({confidence:.0%})", _default_progress.get("percent", 50))
        return lang_code, lang_name
    except Exception as e:
        print(f"\u26a0\ufe0f Language detection failed: {e}. Defaulting to English")
        return 'en', 'English'


def check_cancel():
    """Abort processing if user requested cancellation for current session."""
    session_id = get_current_session()
    if is_cancelled(session_id):
        update_progress("\u274c Processing canceled.", 0)
        raise Exception("Processing canceled by user.")


def reset_cancel(session_id=None):
    """Reset cancel flag and prepare session for new processing."""
    if session_id:
        set_current_session(session_id)
        clear_session(session_id)  # Clear any previous state
        set_progress_status(session_id, "Starting...", 0)


def apply_noise_reduction(audio, reduction_factor=0.8):
    """Apply noise reduction by normalizing and lowering quiet parts of audio."""
    try:
        # Normalize audio first
        normalized = normalize(audio)
        # Apply slight compression to reduce noise
        reduced = normalized.apply_gain_dbfs(-2)
        return reduced
    except Exception as e:
        print(f"⚠️ Warning: Noise reduction failed: {e}. Continuing with original audio.")
        return audio


def extract_audio(video_path, output_path):
    update_progress("🎧 Extracting and processing audio...", 10)
    check_cancel()
    clip = None
    try:
        clip = mp.VideoFileClip(video_path)
        if clip.audio is None:
            raise ValueError("Video file has no audio track")
        
        # Extract audio format
        update_progress("🎧 Converting audio format...", 15)
        clip.audio.write_audiofile(
            output_path, 
            verbose=False, 
            logger=None
        )
        
        # Ensure file is written before processing
        import time
        for _ in range(5):
            if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                break
            time.sleep(0.2)
        
        if not os.path.exists(output_path):
            raise FileNotFoundError(f"Audio file was not created: {output_path}")
        
        # Apply noise reduction for cleaner transcription
        update_progress("🔧 Applying noise reduction...", 20)
        try:
            audio = AudioSegment.from_wav(output_path)
            processed_audio = apply_noise_reduction(audio)
            processed_audio.export(output_path, format="wav")
        except Exception as audio_error:
            print(f"⚠️ Warning: Could not process audio: {audio_error}. Using original.")
        
        update_progress("✅ Audio processing completed.", 25)
        return True
    except Exception as e:
        print(f"❌ Error extracting audio: {e}")
        return False
    finally:
        if clip is not None:
            clip.close()
            if clip.audio is not None:
                clip.audio.close()


def transcribe_audio(audio_path, language='en'):
    """Transcribe audio using Faster-Whisper with HIGH ACCURACY settings.
    
    Args:
        audio_path: Path to audio file
        language: Language code (e.g., 'en', 'es', 'fr', etc.)
    """
    import tempfile
    
    update_progress(f"\ud83d\udde3\ufe0f Transcribing audio ({LANGUAGE_NAMES.get(language, 'English')})...", 40)
    check_cancel()

    global _whisper_model
    try:
        # Thread-safe model initialization
        with _model_lock:
            if _whisper_model is None:
                device = get_device()
                compute_type = get_compute_type(device)
                # Use model size from env (default: base for accuracy)
                model_size = os.environ.get('WHISPER_MODEL_SIZE', 'base')
                _whisper_model = WhisperModel(model_size, device=device, compute_type=compute_type)
                print(f"\u2705 Whisper-{model_size} on {device} ({compute_type}) - High Accuracy mode")
        model = _whisper_model

        # Load audio
        audio = AudioSegment.from_wav(audio_path)
        # 3-minute chunks for better context
        chunk_length_ms = 3 * 60 * 1000
        chunks = [audio[i:i + chunk_length_ms] for i in range(0, len(audio), chunk_length_ms)]

        full_text = ""
        total_chunks = len(chunks)
        temp_files = []  # Track temp files for cleanup
        
        try:
            for i, chunk in enumerate(chunks, start=1):
                check_cancel()

                # Use tempfile for safer temp file handling
                temp_fd, temp_chunk_path = tempfile.mkstemp(suffix='.wav', prefix=f'chunk_{i}_')
                temp_files.append(temp_chunk_path)
                os.close(temp_fd)  # Close fd, we'll use path
                
                try:
                    chunk.export(temp_chunk_path, format="wav")
                except Exception as e:
                    print(f"\u274c Error exporting chunk {i}: {e}")
                    continue

                update_progress(f"\ud83d\udde3\ufe0f Chunk {i}/{total_chunks}...", 40 + int((i / total_chunks) * 20))

                try:
                    # HIGH ACCURACY settings
                    segments, info = model.transcribe(
                        temp_chunk_path, 
                        language=language,
                        beam_size=5,        # Larger beam = better accuracy
                        vad_filter=True,    # Enable VAD for quality
                        vad_parameters=dict(min_silence_duration_ms=300),
                        temperature=0.0     # Deterministic
                    )
                    chunk_text = " ".join([seg.text.strip() for seg in segments if seg.text.strip()])
                    if chunk_text:
                        full_text += chunk_text + " "
                        print(f"\u2705 Chunk {i}: {len(chunk_text.split())} words")
                except Exception as e:
                    print(f"\u274c Error transcribing chunk {i}: {e}")
        finally:
            # Cleanup all temp files
            for temp_path in temp_files:
                try:
                    if os.path.exists(temp_path):
                        os.remove(temp_path)
                except Exception as cleanup_error:
                    print(f"\u26a0\ufe0f Could not remove temp: {cleanup_error}")

        if not full_text.strip():
            update_progress("⚠️ Transcription failed (no text).", 55)
            return ""

        cleaned_text = " ".join(full_text.split())
        print(f"📊 Total: {len(cleaned_text.split())} words")
        update_progress("✅ Transcription completed.", 60)
        return cleaned_text.strip()

    except Exception as e:
        print(f"❌ Transcription error: {e}")
        import traceback
        traceback.print_exc()
        return ""


def summarize_text(text):
    update_progress("\ud83e\udde0 Summarizing text...", 65)
    check_cancel()
    global _summarizer
    try:
        # Validate input
        if not text or not text.strip():
            raise ValueError("Empty text cannot be summarized")
        
        # If text is too short, just return as-is
        words = text.split()
        if len(words) < 50:
            update_progress("\u2705 Summary generation completed.", 85)
            return text
        
        # Thread-safe model initialization
        with _model_lock:
            if _summarizer is None:
                try:
                    print("\ud83d\udd04 Loading BART summarization model...")
                    device = 0 if torch.cuda.is_available() else -1  # Use GPU if available
                    tokenizer = AutoTokenizer.from_pretrained("facebook/bart-large-cnn")
                    model = AutoModelForSeq2SeqLM.from_pretrained("facebook/bart-large-cnn")
                    if device == 0:
                        model = model.to("cuda")
                    _summarizer = {"model": model, "tokenizer": tokenizer, "device": device}
                    print("\u2705 Summarizer model loaded successfully")
                except Exception as model_error:
                    print(f"\u274c Failed to load summarizer model: {model_error}")
                    raise
        
        model = _summarizer["model"]
        tokenizer = _summarizer["tokenizer"]
        device = _summarizer["device"]
        
        # Optimize: Use chunks of appropriate size
        max_chunk_length = 1024  # Max tokens for BART input
        text_chunks = []
        
        # Split text into sentences for better chunking
        sentences = text.split('. ')
        current_chunk = ""
        
        for sentence in sentences:
            if len((current_chunk + " " + sentence).split()) < 200:
                current_chunk += " " + sentence if current_chunk else sentence
            else:
                if current_chunk:
                    text_chunks.append(current_chunk + ".")
                current_chunk = sentence
        if current_chunk:
            text_chunks.append(current_chunk + ".")
        
        summaries = []
        
        for idx, chunk in enumerate(text_chunks):
            check_cancel()
            if not chunk.strip() or len(chunk.split()) < 20:
                continue
            
            try:
                # Tokenize and summarize
                inputs = tokenizer.encode(chunk, return_tensors="pt", max_length=1024, truncation=True)
                if device == 0:
                    inputs = inputs.to("cuda")
                
                summary_ids = model.generate(
                    inputs, 
                    max_length=150, 
                    min_length=30, 
                    num_beams=2,
                    early_stopping=True
                )
                
                summary_text = tokenizer.decode(summary_ids[0], skip_special_tokens=True)
                if summary_text.strip():
                    summaries.append(summary_text)
            except Exception as chunk_error:
                print(f"⚠️ Error summarizing chunk {idx + 1}: {chunk_error}")
                # Use extractive approach as fallback
                words_in_chunk = chunk.split()
                if len(words_in_chunk) > 50:
                    summaries.append(" ".join(words_in_chunk[:len(words_in_chunk)//3]))
        
        final_summary = " ".join(summaries)
        if not final_summary.strip():
            # If all chunks failed, return first third of text
            final_summary = " ".join(text.split()[:len(text.split())//3])
        
        update_progress("✅ Summary generation completed.", 85)
        return final_summary
    except Exception as e:
        print(f"❌ Error during summarization: {e}")
        import traceback
        traceback.print_exc()
        raise  # Re-raise so app.py can handle it properly


def text_to_speech(text, output_audio_path, language='en', speed_factor=1.15):
    """
    Generate speech from text with adjustable speed and language support.
    
    Args:
        text: Text to convert to speech
        output_audio_path: Path to save the audio file
        language: Language code (e.g., 'en', 'es', 'fr', 'hi', 'ta', 'ja', 'zh', 'ko', 'ar', 'ru', 'de', 'it', 'pt')
        speed_factor: 1.0 = normal, 1.15 = slightly faster but still clear
    """
    update_progress("🔊 Generating summary speech...", 90)
    check_cancel()
    
    # Map language codes to gTTS supported codes
    GTTS_LANG_MAP = {
        'en': 'en',
        'es': 'es',
        'fr': 'fr',
        'de': 'de',
        'it': 'it',
        'pt': 'pt',
        'ru': 'ru',
        'ja': 'ja',
        'zh': 'zh-CN',  # Chinese Simplified
        'hi': 'hi',
        'ar': 'ar',
        'ko': 'ko',
        'ta': 'ta',
    }
    
    # Get the gTTS language code
    gtts_lang = GTTS_LANG_MAP.get(language, 'en')
    
    try:
        if not text or not text.strip():
            raise ValueError("No text to speak")
        
        import tempfile
        
        try:
            # Generate initial audio with gTTS in specified language
            print(f"🔊 Generating TTS in language: {gtts_lang}")
            tts = gTTS(text=text, lang=gtts_lang, slow=False)
            
            # Try to use pydub for speed adjustment if ffmpeg is available
            try:
                # Save to temporary file first
                with tempfile.NamedTemporaryFile(delete=False, suffix='.mp3') as temp_file:
                    temp_path = temp_file.name
                tts.save(temp_path)
                
                # Speed up the audio using pydub for clearer, faster speech
                audio = AudioSegment.from_mp3(temp_path)
                
                # Apply speed adjustment (1.15x is fast but still understandable)
                if speed_factor != 1.0:
                    # Speed up by changing frame rate, then convert back
                    faster_audio = audio._spawn(audio.raw_data, overrides={
                        "frame_rate": int(audio.frame_rate * speed_factor)
                    }).set_frame_rate(audio.frame_rate)
                    faster_audio.export(output_audio_path, format="mp3", bitrate="192k")
                else:
                    audio.export(output_audio_path, format="mp3", bitrate="192k")
                
                # Clean up temp file
                try:
                    os.remove(temp_path)
                except:
                    pass
                    
            except Exception as pydub_error:
                # Pydub/ffmpeg not available - save directly without speed adjustment
                print(f"⚠️ Pydub not available ({pydub_error}), saving audio directly without speed adjustment")
                tts.save(output_audio_path)
            
            print(f"✅ Audio generated successfully: {output_audio_path}")
                
        except Exception as tts_error:
            print(f"⚠️ gTTS failed: {tts_error}")
            raise tts_error
        
        update_progress("✅ Summary speech generated.", 100)
    except Exception as e:
        print(f"❌ Error generating summary audio: {e}")
        import traceback
        traceback.print_exc()
        raise  # Re-raise exception so caller knows it failed


def translate_text(text, target_language='es'):
    """Translate text to target language using Google Translate.
    
    Args:
        text: Text to translate
        target_language: Target language code (e.g., 'es' for Spanish, 'fr' for French)
    
    Returns:
        Translated text or original if translation fails
    """
    import asyncio
    
    async def _translate_async():
        try:
            if not text or not text.strip():
                return ""
            
            # Language code mapping
            LANGUAGE_CODES = {
                'en': 'English',
                'es': 'Spanish',
                'fr': 'French',
                'de': 'German',
                'it': 'Italian',
                'pt': 'Portuguese',
                'ru': 'Russian',
                'ja': 'Japanese',
                'zh': 'Chinese',
                'hi': 'Hindi',
                'ar': 'Arabic',
                'ko': 'Korean',
                'ta': 'Tamil',
            }
            
            target_lang = target_language if target_language in LANGUAGE_CODES else 'es'
            
            translator = Translator()
            
            # Split text into chunks to avoid API limits
            chunks = []
            words = text.split()
            current_chunk = []
            current_length = 0
            
            for word in words:
                word_length = len(word) + 1
                if current_length + word_length > 400:
                    if current_chunk:
                        chunks.append(' '.join(current_chunk))
                    current_chunk = [word]
                    current_length = word_length
                else:
                    current_chunk.append(word)
                    current_length += word_length
            
            if current_chunk:
                chunks.append(' '.join(current_chunk))
            
            # Translate each chunk
            translated_chunks = []
            for chunk in chunks:
                try:
                    # googletrans 4.0.2: pass language code as positional argument
                    translation = await translator.translate(chunk, target_lang)
                    if hasattr(translation, 'text'):
                        translated_chunks.append(translation.text)
                    else:
                        translated_chunks.append(str(translation))
                except Exception as e:
                    print(f"⚠️ Translation chunk failed: {e}. Using original chunk.")
                    translated_chunks.append(chunk)
            
            translated_text = ' '.join(translated_chunks)
            print(f"✅ Translation to {LANGUAGE_CODES.get(target_lang, 'Unknown')} completed")
            
            return translated_text
        
        except Exception as e:
            print(f"❌ Translation error: {e}")
            return text
    
    # Run async function
    try:
        loop = asyncio.get_event_loop()
        if loop.is_running():
            # Create new event loop in thread if already running
            import concurrent.futures
            import threading
            result = [None]
            def run_in_thread():
                result[0] = asyncio.run(_translate_async())
            thread = threading.Thread(target=run_in_thread)
            thread.start()
            thread.join()
            return result[0]
        else:
            return asyncio.run(_translate_async())
    except:
        # Fallback: try to run directly
        return asyncio.run(_translate_async())


def export_to_pdf(video_name, transcript, summary, metadata=None):
    """Export transcript and summary to PDF format with Unicode support for multiple languages"""
    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from datetime import datetime
        import os
        
        # Get the target language from metadata
        target_language = metadata.get('language', 'en') if metadata else 'en'
        print(f"📄 Generating PDF for language: {target_language}")
        
        # Language-specific font mapping (fonts in static/fonts/)
        # Each language maps to (font_path, font_name)
        LANGUAGE_FONTS = {
            # Indian Languages
            'ta': ('static/fonts/NotoSansTamil-Regular.ttf', 'NotoSansTamil'),           # Tamil
            'hi': ('static/fonts/NotoSansDevanagari-Regular.ttf', 'NotoSansDevanagari'), # Hindi
            # Arabic
            'ar': ('static/fonts/NotoSansArabic-Regular.ttf', 'NotoSansArabic'),         # Arabic
            # CJK Languages - dedicated fonts for each
            'ja': ('static/fonts/NotoSansJP-Regular.ttf', 'NotoSansJP'),                 # Japanese
            'zh': ('static/fonts/NotoSansSC-Regular.ttf', 'NotoSansSC'),                 # Chinese (Simplified)
            'ko': ('static/fonts/NotoSansKR-Regular.ttf', 'NotoSansKR'),                 # Korean
            # Western/Cyrillic Languages (English, Spanish, French, German, Italian, Portuguese, Russian)
            'en': ('static/fonts/NotoSans-Regular.ttf', 'NotoSans'),                      # English
            'es': ('static/fonts/NotoSans-Regular.ttf', 'NotoSans'),                      # Spanish
            'fr': ('static/fonts/NotoSans-Regular.ttf', 'NotoSans'),                      # French
            'de': ('static/fonts/NotoSans-Regular.ttf', 'NotoSans'),                      # German
            'it': ('static/fonts/NotoSans-Regular.ttf', 'NotoSans'),                      # Italian
            'pt': ('static/fonts/NotoSans-Regular.ttf', 'NotoSans'),                      # Portuguese
            'ru': ('static/fonts/NotoSans-Regular.ttf', 'NotoSans'),                      # Russian (Cyrillic supported by NotoSans)
        }
        
        # Default font for any unlisted languages
        DEFAULT_FONT = ('static/fonts/NotoSans-Regular.ttf', 'NotoSans')
        
        # Select the appropriate font based on language
        font_path, font_name = LANGUAGE_FONTS.get(target_language, DEFAULT_FONT)
        
        # Register fonts - English font for metadata/labels, language font for content
        english_font_name = 'Helvetica'  # Fallback for English labels
        content_font_name = 'Helvetica'  # Fallback for translated content
        
        def safe_register_font(name, path):
            """Register font if not already registered"""
            try:
                # Check if font already registered
                registered_fonts = pdfmetrics.getRegisteredFontNames()
                if name in registered_fonts:
                    return name
                if os.path.exists(path):
                    pdfmetrics.registerFont(TTFont(name, path))
                    return name
            except Exception as e:
                print(f"Font registration warning for {name}: {e}")
            return None
        
        try:
            # Register NotoSans for English labels/metadata
            english_font_path = 'static/fonts/NotoSans-Regular.ttf'
            result = safe_register_font('NotoSans', english_font_path)
            if result:
                english_font_name = result
                print(f"✅ Using English font: NotoSans for labels/metadata")
            
            # Register language-specific font for content
            result = safe_register_font(font_name, font_path)
            if result:
                content_font_name = result
                print(f"✅ Using content font: {font_name} for {target_language}")
            else:
                content_font_name = english_font_name  # Fall back to English font
                
        except Exception as font_error:
            print(f"⚠️ Font registration error: {font_error}")
            # Try reportlab's built-in DejaVuSans
            try:
                import reportlab
                reportlab_path = os.path.dirname(reportlab.__file__)
                dejavu_path = os.path.join(reportlab_path, 'fonts', 'DejaVuSans.ttf')
                result = safe_register_font('DejaVuSans', dejavu_path)
                if result:
                    english_font_name = result
                    content_font_name = result
            except:
                pass
        
        # Create PDF with professional margins
        # Build filename: Report_Language_Filename.pdf
        language_name = metadata.get('language_name', 'English') if metadata else 'English'
        # Clean video name - remove special characters and timestamp suffix
        clean_video_name = re.sub(r'_\d{10,}$', '', video_name)  # Remove trailing timestamp
        clean_video_name = ''.join(c for c in clean_video_name if c.isalnum() or c in ' _-').strip().replace(' ', '_')
        if not clean_video_name:
            clean_video_name = 'VideoReport'
        pdf_filename = f"Report_{language_name}_{clean_video_name}.pdf"
        pdf_path = f"static/exports/{pdf_filename}"
        os.makedirs('static/exports', exist_ok=True)
        
        # ===== WATERMARK FUNCTION =====
        def add_watermark(canvas, doc):
            """Add VidSummify watermark logo to every page - centered, light grey, transparent"""
            canvas.saveState()
            
            # Get page dimensions
            page_width, page_height = letter
            
            # Set watermark properties - light grey color for transparency effect
            canvas.setFillColorRGB(0.85, 0.85, 0.85)  # Light grey (85% grey)
            
            # Set font for watermark text
            canvas.setFont('Helvetica-Bold', 60)
            
            # Calculate center position
            text = "VidSummify"
            text_width = canvas.stringWidth(text, 'Helvetica-Bold', 60)
            x = (page_width - text_width) / 2
            y = page_height / 2
            
            # Rotate and draw the watermark diagonally for professional look
            canvas.translate(page_width / 2, page_height / 2)
            canvas.rotate(45)
            canvas.drawCentredString(0, 0, text)
            
            # Add smaller tagline below
            canvas.setFont('Helvetica', 20)
            canvas.setFillColorRGB(0.88, 0.88, 0.88)  # Slightly lighter grey
            canvas.drawCentredString(0, -40, "AI Video Summarization")
            
            canvas.restoreState()
        
        # Professional document margins (left, right, top, bottom)
        doc = SimpleDocTemplate(
            pdf_path, 
            pagesize=letter,
            leftMargin=0.75*inch,
            rightMargin=0.75*inch,
            topMargin=0.75*inch,
            bottomMargin=0.75*inch
        )
        story = []
        styles = getSampleStyleSheet()
        
        # Import additional reportlab components for professional styling
        from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_LEFT
        from reportlab.platypus import HRFlowable, Table, TableStyle
        from reportlab.lib import colors
        
        # Custom styles - English font for metadata/labels, content font for translated text
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontName=english_font_name,
            fontSize=22,
            textColor=colors.HexColor('#667eea'),
            spaceAfter=8,
            alignment=TA_CENTER,
            leading=28
        )
        
        subtitle_style = ParagraphStyle(
            'CustomSubtitle',
            parent=styles['Normal'],
            fontName=english_font_name,
            fontSize=10,
            textColor=colors.HexColor('#718096'),
            alignment=TA_CENTER,
            spaceAfter=20
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontName=english_font_name,
            fontSize=14,
            textColor=colors.HexColor('#1a202c'),
            spaceBefore=16,
            spaceAfter=10,
            borderPadding=4,
            leftIndent=0
        )
        
        # Content style with JUSTIFIED alignment for clean edges
        content_style = ParagraphStyle(
            'CustomContent',
            parent=styles['Normal'],
            fontName=content_font_name,
            fontSize=11,
            leading=18,  # 1.6x line height for readability
            textColor=colors.HexColor('#2d3748'),
            alignment=TA_JUSTIFY,  # Justified text like MS Word
            firstLineIndent=0,
            spaceBefore=4,
            spaceAfter=8
        )
        
        meta_label_style = ParagraphStyle(
            'MetaLabel',
            parent=styles['Normal'],
            fontName=english_font_name,
            fontSize=9,
            textColor=colors.HexColor('#4a5568'),
            alignment=TA_LEFT
        )
        
        meta_value_style = ParagraphStyle(
            'MetaValue',
            parent=styles['Normal'],
            fontName=english_font_name,
            fontSize=9,
            textColor=colors.HexColor('#1a202c'),
            alignment=TA_LEFT
        )
        
        # ===== DOCUMENT HEADER =====
        story.append(Paragraph(f"VidSummify Report", title_style))
        story.append(Paragraph(video_name, subtitle_style))
        
        # Add decorative line under title
        story.append(HRFlowable(width="100%", thickness=2, color=colors.HexColor('#667eea'), spaceBefore=5, spaceAfter=15))
        
        if metadata:
            # Get language name for display
            display_language = metadata.get('language_name', metadata.get('language', 'English'))
            LANG_DISPLAY_NAMES = {
                'en': 'English', 'es': 'Spanish', 'fr': 'French', 'de': 'German',
                'it': 'Italian', 'pt': 'Portuguese', 'ru': 'Russian', 'ja': 'Japanese',
                'zh': 'Chinese', 'hi': 'Hindi', 'ar': 'Arabic', 'ko': 'Korean', 'ta': 'Tamil',
            }
            if display_language in LANG_DISPLAY_NAMES:
                display_language = LANG_DISPLAY_NAMES[display_language]
            
            # Calculate word counts
            transcript_words = len(transcript.split()) if transcript else 0
            summary_words = len(summary.split()) if summary else 0
            compression = round((1 - summary_words / transcript_words) * 100, 1) if transcript_words > 0 else 0
            
            # Create metadata table for professional layout
            meta_data = [
                [Paragraph("<b>Generated</b>", meta_label_style), 
                 Paragraph(datetime.now().strftime('%Y-%m-%d %H:%M:%S'), meta_value_style),
                 Paragraph("<b>Language</b>", meta_label_style),
                 Paragraph(display_language, meta_value_style)],
                [Paragraph("<b>Processing Time</b>", meta_label_style), 
                 Paragraph(str(metadata.get('duration', 'N/A')), meta_value_style),
                 Paragraph("<b>Compression</b>", meta_label_style),
                 Paragraph(f"{compression}%", meta_value_style)],
                [Paragraph("<b>Transcript Words</b>", meta_label_style), 
                 Paragraph(f"{transcript_words:,}", meta_value_style),
                 Paragraph("<b>Summary Words</b>", meta_label_style),
                 Paragraph(f"{summary_words:,}", meta_value_style)],
            ]
            
            meta_table = Table(meta_data, colWidths=[1.4*inch, 1.8*inch, 1.4*inch, 1.8*inch])
            meta_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#f8fafc')),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#e2e8f0')),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ('LEFTPADDING', (0, 0), (-1, -1), 10),
                ('RIGHTPADDING', (0, 0), (-1, -1), 10),
            ]))
            story.append(meta_table)
            story.append(Spacer(1, 0.4*inch))
        
        # ===== TRANSCRIPT SECTION =====
        # Section header with decorative underline
        story.append(Paragraph("Transcript", heading_style))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#e2e8f0'), spaceBefore=2, spaceAfter=12))
        
        # Normalize Unicode text for proper rendering of complex scripts (Tamil, Hindi, Arabic, etc.)
        normalized_transcript = unicodedata.normalize('NFC', transcript)
        
        # Handle Arabic text with proper reshaping and RTL
        if target_language == 'ar':
            try:
                import arabic_reshaper
                from bidi.algorithm import get_display
                normalized_transcript = get_display(arabic_reshaper.reshape(normalized_transcript))
            except ImportError:
                pass  # Use original if libraries not available
        
        # Escape special characters and handle line breaks
        safe_transcript = normalized_transcript.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('\n', '<br/>')
        story.append(Paragraph(safe_transcript, content_style))
        story.append(PageBreak())
        
        # ===== SUMMARY SECTION =====
        story.append(Paragraph("Summary", heading_style))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#e2e8f0'), spaceBefore=2, spaceAfter=12))
        
        # Normalize Unicode text for summary as well
        normalized_summary = unicodedata.normalize('NFC', summary)
        
        # Handle Arabic text with proper reshaping and RTL
        if target_language == 'ar':
            try:
                import arabic_reshaper
                from bidi.algorithm import get_display
                normalized_summary = get_display(arabic_reshaper.reshape(normalized_summary))
            except ImportError:
                pass  # Use original if libraries not available
        
        safe_summary = normalized_summary.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('\n', '<br/>')
        story.append(Paragraph(safe_summary, content_style))
        
        # Add footer spacer
        story.append(Spacer(1, 0.5*inch))
        
        # Add footer with branding
        footer_style = ParagraphStyle(
            'Footer',
            parent=styles['Normal'],
            fontName=english_font_name,
            fontSize=8,
            textColor=colors.HexColor('#a0aec0'),
            alignment=TA_CENTER
        )
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#e2e8f0'), spaceBefore=10, spaceAfter=10))
        story.append(Paragraph("Generated by VidSummify - AI-Powered Video Summarization", footer_style))
        
        # Build PDF with watermark on every page
        doc.build(story, onFirstPage=add_watermark, onLaterPages=add_watermark)
        return pdf_path
    except Exception as e:
        print(f"❌ Error exporting to PDF: {e}")
        import traceback
        traceback.print_exc()
        raise


def export_to_docx(video_name, transcript, summary, metadata=None):
    """Export transcript and summary to DOCX format"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from datetime import datetime
        
        # Create Document
        doc = Document()
        
        # Add title
        title = doc.add_heading(f'VidSummify Report: {video_name}', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        if metadata:
            doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph(f"Duration: {metadata.get('duration', 'N/A')}")
            doc.add_paragraph(f"Language: {metadata.get('language', 'English')}")
            doc.add_paragraph()
        
        # Add transcript
        doc.add_heading('Transcript', level=2)
        doc.add_paragraph(transcript)
        doc.add_page_break()
        
        # Add summary
        doc.add_heading('Summary', level=2)
        doc.add_paragraph(summary)
        
        # Save document
        # Build filename: Report_Language_Filename.docx
        language_name = metadata.get('language_name', 'English') if metadata else 'English'
        clean_video_name = ''.join(c for c in video_name if c.isalnum() or c in ' _-').strip().replace(' ', '_')
        docx_filename = f"Report_{language_name}_{clean_video_name}.docx"
        docx_path = f"static/exports/{docx_filename}"
        doc.save(docx_path)
        return docx_path
    except Exception as e:
        print(f"❌ Error exporting to DOCX: {e}")
        raise


def export_to_json(video_name, transcript, summary, metadata=None):
    """Export to JSON format"""
    import json
    from datetime import datetime
    
    try:
        data = {
            'video_name': video_name,
            'generated_at': datetime.now().isoformat(),
            'transcript': transcript,
            'summary': summary,
            'metadata': metadata or {}
        }
        
        # Build filename: Report_Language_Filename.json
        language_name = metadata.get('language_name', 'English') if metadata else 'English'
        clean_video_name = ''.join(c for c in video_name if c.isalnum() or c in ' _-').strip().replace(' ', '_')
        json_filename = f"Report_{language_name}_{clean_video_name}.json"
        json_path = f"static/exports/{json_filename}"
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        
        return json_path
    except Exception as e:
        print(f"❌ Error exporting to JSON: {e}")
        raise


def export_to_srt(video_name, transcript, summary=None):
    """Export transcript as SRT subtitle format"""
    try:
        lines = transcript.split('\n')
        srt_content = ""
        chunk_size = 10  # Words per subtitle
        words = transcript.split()
        
        index = 1
        for i in range(0, len(words), chunk_size):
            chunk = ' '.join(words[i:i+chunk_size])
            # Simple timing (assuming ~150 words per minute = 2.5 words per second)
            start_time = f"00:00:{int(i/(chunk_size*2.5)):02d},000"
            end_time = f"00:00:{int((i+chunk_size)/(chunk_size*2.5)):02d},000"
            srt_content += f"{index}\n{start_time} --> {end_time}\n{chunk}\n\n"
            index += 1
        
        srt_path = f"static/exports/{video_name}_export.srt"
        with open(srt_path, 'w', encoding='utf-8') as f:
            f.write(srt_content)
        
        return srt_path
    except Exception as e:
        print(f"❌ Error exporting to SRT: {e}")
        raise

